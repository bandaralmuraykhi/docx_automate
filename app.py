import os
import tempfile
import uuid
import hashlib
from flask import Flask, session, render_template, request, redirect, jsonify, send_file, abort, url_for, flash
from flask_sqlalchemy import SQLAlchemy
from flask_login import LoginManager, UserMixin, login_user, logout_user, login_required, current_user
from flask_mail import Mail, Message
from flask_wtf import FlaskForm
from wtforms import StringField, PasswordField, SubmitField, SelectMultipleField
from wtforms.validators import DataRequired, Length, Email, EqualTo
from docx import Document
from itsdangerous import URLSafeTimedSerializer as Serializer

app = Flask(__name__)
app.static_folder = ''
# Access the environment variable `export FLASK_SECRET_KEY=your_secret_key_here`
app.config['SECRET_KEY'] = 'FLASK_SECRET_KEY'
# app.config['SECRET_KEY'] = os.environ.get('FLASK_SECRET_KEY')
basedir = os.path.abspath(os.path.dirname(__file__))
app.config['SQLALCHEMY_DATABASE_URI'] = 'sqlite:///' + os.path.join(basedir, 'instance', 'database.db')
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False
db = SQLAlchemy(app)
login_manager = LoginManager(app)
login_manager.login_view = 'login'

# Configuration for Flask-Mail
app.config.from_pyfile('config.cfg')
mail = Mail(app)

UPLOAD_FOLDER = 'uploads'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)

def generate_salt():
    # Generate a 16-byte salt
    return os.urandom(16)

def hash_password(password, salt):
    # Hash a password with the provided salt
    pwdhash = hashlib.pbkdf2_hmac('sha256', password.encode('utf-8'), salt, 100000)
    return pwdhash

@login_manager.user_loader
def load_user(user_id):
    return User.query.get(int(user_id))

class User(UserMixin, db.Model):
    id = db.Column(db.Integer, primary_key=True)
    email = db.Column(db.String(100), unique=True, nullable=False)
    password = db.Column(db.String(256), nullable=False)
    salt = db.Column(db.LargeBinary(16), nullable=False)
    confirmed = db.Column(db.Boolean, default=False)

    def generate_confirmation_token(self, expiration=3600):
        s = Serializer(app.config['SECRET_KEY'])
        return s.dumps({'confirm': self.id})

    def confirm(self, token):
        s = Serializer(app.config['SECRET_KEY'])
        try:
            data = s.loads(token, max_age=3600)
        except:
            return False
        if data.get('confirm') != self.id:
            return False
        self.confirmed = True
        db.session.commit()
        return True

    def generate_reset_token(self, expiration=3600):
        s = Serializer(app.config['SECRET_KEY'])
        return s.dumps({'reset': self.id})

    @staticmethod
    def reset_password(token, new_password):
        s = Serializer(app.config['SECRET_KEY'])
        try:
            data = s.loads(token, max_age=3600)
        except:
            return False
        user = User.query.get(data.get('reset'))
        if user is None:
            return False
        user.password = new_password
        db.session.commit()
        return True
    
    @staticmethod
    def verify_reset_token(token):
        s = Serializer(app.config['SECRET_KEY'])
        try:
            data = s.loads(token)
        except:
            return None
        return User.query.get(data.get('reset'))

class FormResponse(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    form_id = db.Column(db.String(100), nullable=False)
    form_name = db.Column(db.String(100), nullable=False)
    form_link = db.Column(db.String(200), nullable=False)
    selected_cells = db.Column(db.String(500), nullable=False)
    unique_filename = db.Column(db.String(100), nullable=False)
    user_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    user = db.relationship('User', backref='form_responses')
    send_email = db.Column(db.Boolean, default=False)
    allow_download = db.Column(db.Boolean, default=True)

class SignupForm(FlaskForm):
    email = StringField('Email', validators=[DataRequired(), Email()])
    password = PasswordField('Password', validators=[DataRequired(), Length(min=6)])
    confirm_password = PasswordField('Confirm Password', validators=[DataRequired(), EqualTo('password')])
    submit = SubmitField('Sign Up')

class LoginForm(FlaskForm):
    email = StringField('Email', validators=[DataRequired(), Email()])
    password = PasswordField('Password', validators=[DataRequired()])
    submit = SubmitField('Log In')

class UpdateFormForm(FlaskForm):
    form_name = StringField('Form Name', validators=[DataRequired()])
    selected_cells = SelectMultipleField('Selected Cells', coerce=str)
    submit = SubmitField('Update')

class ResetPasswordRequestForm(FlaskForm):
    email = StringField('Email', validators=[DataRequired(), Email()])
    submit = SubmitField('Request Password Reset')

class ResendConfirmationForm(FlaskForm):
    submit = SubmitField('Resend Confirmation Email')

class ResetPasswordForm(FlaskForm):
    password = PasswordField('New Password', validators=[DataRequired(), Length(min=6)])
    confirm_password = PasswordField('Confirm Password', validators=[DataRequired(), EqualTo('password')])
    submit = SubmitField('Reset Password')

@app.route('/set_theme', methods=['POST'])
def set_theme():
    theme = request.json.get('theme')
    if theme in ['light', 'dark']:
        session['theme'] = theme
        return jsonify({'success': True})
    else:
        return jsonify({'success': False})

@app.route('/')
def home():
    return render_template('home.html', theme=session.get('theme', 'light'))

@app.route('/upload')
@login_required
def upload():
    return render_template('upload.html', theme=session.get('theme', 'light'))

@app.route('/signup', methods=['GET', 'POST'])
def signup():
    form = SignupForm()
    if form.validate_on_submit():
        try:
            existing_user = User.query.filter_by(email=form.email.data).first()
            if existing_user:
                flash('Email address already exists. Please log in.', 'warning')
                return redirect(url_for('login'))

            salt = generate_salt()
            hashed_password = hash_password(form.password.data, salt)
            new_user = User(email=form.email.data, password=hashed_password, salt=salt)
            db.session.add(new_user)
            db.session.commit()

            token = new_user.generate_confirmation_token()
            send_email(new_user.email, 'Confirm Your Email', 'confirm_email', user=new_user, token=token)
            flash('A confirmation email has been sent to you.', 'success')
            return redirect(url_for('unconfirmed', email=new_user.email))
        except Exception as e:
            db.session.rollback()
            flash('An error occurred during signup. Please try again.', 'danger')
            app.logger.error(f"Signup error: {str(e)}")

    return render_template('signup.html', form=form, theme=session.get('theme', 'light'))

@app.route('/login', methods=['GET', 'POST'])
def login():
    form = LoginForm()
    if form.validate_on_submit():
        user = User.query.filter_by(email=form.email.data).first()
        if user:
            hashed_password = hash_password(form.password.data, user.salt)
            if hashed_password == user.password:
                if user.confirmed:
                    login_user(user)
                    # flash('Login successful.', 'success')
                    return redirect(url_for('home'))
                else:
                    flash('Please confirm your email address to access your account.', 'warning')
                    return redirect(url_for('unconfirmed', email=user.email))
            else:
                flash('Invalid email or password. Please try again.', 'danger')
        else:
            flash('Invalid email or password. Please try again.', 'danger')

    return render_template('login.html', form=form, theme=session.get('theme', 'light'))

@app.route('/logout')
@login_required
def logout():
    try:
        logout_user()
        flash('You have been logged out.', 'success')
        return redirect(url_for('login'))
    except Exception as e:
        flash('An error occurred during logout.', 'danger')
        app.logger.error(f"Logout error: {str(e)}")
        return redirect(url_for('home'))

@app.route('/confirm/<token>')
def confirm(token):
    try:
        user = User.query.filter_by(confirmed=False).first()
        if user and user.confirm(token):
            flash('Your email has been confirmed.', 'success')
        else:
            flash('The confirmation link is invalid or has expired.', 'danger')
    except Exception as e:
        flash('An error occurred during email confirmation.', 'danger')
        app.logger.error(f"Email confirmation error: {str(e)}")

    return redirect(url_for('login'))

@app.route('/unconfirmed', methods=['GET', 'POST'])
def unconfirmed():
    email = request.args.get('email')
    user = User.query.filter_by(email=email).first()

    if user and user.confirmed:
        return redirect(url_for('upload'))

    form = ResendConfirmationForm()
    if form.validate_on_submit():
        if user:
            try:
                token = user.generate_confirmation_token()
                send_email(user.email, 'Confirm Your Email', 'confirm_email', user=user, token=token)
                flash('A new confirmation email has been sent to you.', 'success')
            except Exception as e:
                flash('An error occurred while sending the confirmation email.', 'danger')
                app.logger.error(f"Confirmation email sending error: {str(e)}")
        else:
            flash('User not found. Please sign up again.', 'danger')

        return redirect(url_for('unconfirmed', email=email))

    flash('Your email address is not confirmed. Please check your inbox and confirm your email.', 'warning')
    return render_template('unconfirmed.html', form=form, theme=session.get('theme', 'light'))

@app.route('/reset_password_request', methods=['GET', 'POST'])
def reset_password_request():
    if current_user.is_authenticated:
        return redirect(url_for('upload'))

    form = ResetPasswordRequestForm()
    if form.validate_on_submit():
        user = User.query.filter_by(email=form.email.data).first()
        if user:
            try:
                token = user.generate_reset_token()
                send_email_reset(user.email, 'Reset Your Password', 'reset_password', user=user, token=token)
                flash('An email with instructions to reset your password has been sent to you.', 'success')
                return redirect(url_for('login'))
            except Exception as e:
                flash('An error occurred while sending the password reset email.', 'danger')
                app.logger.error(f"Password reset email sending error: {str(e)}")
        else:
            flash('No account found with that email address.', 'danger')

    return render_template('reset_password_request.html', form=form, theme=session.get('theme', 'light'))

@app.route('/reset_password/<token>', methods=['GET', 'POST'])
def reset_password(token):
    if current_user.is_authenticated:
        return redirect(url_for('upload'))

    user = User.verify_reset_token(token)
    if not user:
        flash('Invalid or expired reset link.', 'danger')
        return redirect(url_for('reset_password_request'))

    form = ResetPasswordForm()
    if form.validate_on_submit():
        try:
            salt = generate_salt()
            hashed_password = hash_password(form.password.data, salt)
            user.password = hashed_password
            user.salt = salt
            db.session.commit()
            flash('Your password has been updated.', 'success')
            return redirect(url_for('login'))
        except Exception as e:
            db.session.rollback()
            flash('An error occurred while updating your password.', 'danger')
            app.logger.error(f"Password update error: {str(e)}")

    return render_template('reset_password.html', form=form, user=user, theme=session.get('theme', 'light'))

PER_PAGE = 2  # Number of form responses to display per page

@app.route('/dashboard', defaults={'page': 1})
@app.route('/dashboard/page/<int:page>')
@login_required
def dashboard(page):
    try:
        page = request.args.get('page', 1, type=int)
        form_responses = FormResponse.query.filter_by(user_id=current_user.id).paginate(page=page, per_page=PER_PAGE, error_out=False)
        return render_template('dashboard.html', form_responses=form_responses)
    except Exception as e:
        flash('An error occurred while loading the dashboard.', 'danger')
        app.logger.error(f"Dashboard loading error: {str(e)}")
        return redirect(url_for('upload'))

def allowed_file(filename):
    allowed_extensions = ['docx']
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in allowed_extensions

@app.route('/upload', methods=['POST'])
@login_required
def upload_file():
    if 'file' not in request.files:
        flash('No file selected.', 'danger')
        return redirect(url_for('upload'))

    file = request.files['file']
    if file.filename == '':
        flash('No file selected.', 'danger')
        return redirect(url_for('upload'))

    if file and allowed_file(file.filename):
        if file.content_length > 10 * 1024 * 1024:  # Limit file size to 10MB
            flash('File size exceeds the maximum limit of 10MB.', 'danger')
            return redirect(url_for('upload'))

        unique_filename = str(uuid.uuid4()) + '.docx'
        file_path = os.path.join(UPLOAD_FOLDER, unique_filename)
        file.save(file_path)

        try:
            document = Document(file_path)
            table_data = []
            unique_cells = set()
            for table in document.tables:
                for row in table.rows:
                    row_data = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text not in unique_cells:
                            unique_cells.add(cell_text)
                            row_data.append(cell_text)
                    if row_data:
                        table_data.append(row_data)
            return jsonify({'table_data': table_data, 'unique_filename': unique_filename})
        except Exception as e:
            flash('An error occurred while processing the file.', 'danger')
            app.logger.error(f"File processing error: {str(e)}")
            return redirect(url_for('upload'))
    else:
        flash('Invalid file format. Only .docx files are allowed.', 'danger')
        return redirect(url_for('upload'))

@app.route('/create_form', methods=['POST'])
@login_required
def create_form():
    selected_cells = request.json.get('selected_cells', [])
    unique_filename = request.json.get('unique_filename', '')
    send_email = request.json.get('send_email', False)
    allow_download = request.json.get('allow_download', True)

    if not selected_cells or not unique_filename:
        flash('Invalid form data.', 'danger')
        return jsonify({'error': 'Invalid form data.'}), 400

    form_id = str(uuid.uuid4())
    form_url = url_for('fill_form', form_id=form_id, _external=True)
    form_name = f"Form {form_id}"
    form_response = FormResponse(
        form_id=form_id,
        form_name=form_name,
        form_link=form_url,
        selected_cells=','.join(selected_cells),
        unique_filename=unique_filename,
        user_id=current_user.id,
        send_email=send_email,
        allow_download=allow_download
    )

    try:
        db.session.add(form_response)
        db.session.commit()
        return jsonify({
            'form_url': form_url,
            'selected_cells': selected_cells,
            'unique_filename': unique_filename,
            'form_id': form_id,
            'form_name': form_name
        })
    except Exception as e:
        db.session.rollback()
        flash('An error occurred while creating the form.', 'danger')
        app.logger.error(f"Form creation error: {str(e)}")
        return jsonify({'error': 'An error occurred while creating the form.'}), 500

@app.route('/fill_form/<form_id>', methods=['GET', 'POST'])
def fill_form(form_id):
    form_response = FormResponse.query.filter_by(form_id=form_id).first()

    if not form_response:
        flash('Form not found.', 'danger')
        return redirect(url_for('upload'))

    if request.method == 'POST':
        form_data = request.form

        try:
            file_path = os.path.join(UPLOAD_FOLDER, form_response.unique_filename)
            document = Document(file_path)
            for cell_text, value in form_data.items():
                if cell_text != 'unique_filename':
                    for table in document.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                if cell.text.strip() == cell_text:
                                    cell_paragraph = cell.paragraphs[0]
                                    run = cell_paragraph.add_run(f" {value}")
                                    run.font.name = cell_paragraph.runs[0].font.name
                                    run.font.size = cell_paragraph.runs[0].font.size
                                    run.bold = cell_paragraph.runs[0].bold
                                    run.italic = cell_paragraph.runs[0].italic
                                    run.underline = cell_paragraph.runs[0].underline

            temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
            document.save(temp_file.name)
            temp_file.close()

            if form_response.send_email:
                user = User.query.get(form_response.user_id)
                if user:
                    try:
                        msg = Message('Modified File', recipients=[user.email])
                        msg.body = 'Please find the attached modified file.'
                        with open(temp_file.name, 'rb') as f:
                            msg.attach(form_response.form_name + '.docx', 'application/vnd.openxmlformats-officedocument.wordprocessingml.document', f.read())
                        mail.send(msg)
                    except Exception as e:
                        flash('An error occurred while sending the email.', 'danger')
                        app.logger.error(f"Email sending error: {str(e)}")

            if form_response.allow_download:
                return send_file(temp_file.name, as_attachment=True)
            else:
                flash('Form submitted, but download is not allowed.', 'info')
                return redirect(url_for('fill_form', form_id=form_id))
        except Exception as e:
            flash('An error occurred while processing the form.', 'danger')
            app.logger.error(f"Form processing error: {str(e)}")
            return redirect(url_for('fill_form', form_id=form_id))
    else:
        selected_cells = form_response.selected_cells.split(',')
        unique_filename = form_response.unique_filename
        return render_template('fill_form.html', form_id=form_id, selected_cells=selected_cells, unique_filename=unique_filename, theme=session.get('theme', 'light'))

@app.route('/delete_form/<form_id>', methods=['GET'])
@login_required
def delete_form(form_id):
    form_response = FormResponse.query.filter_by(form_id=form_id, user_id=current_user.id).first()

    if not form_response:
        flash('Form not found.', 'danger')
        return redirect(url_for('dashboard'))

    try:
        db.session.delete(form_response)
        db.session.commit()
        flash('Form deleted successfully.', 'success')
    except Exception as e:
        db.session.rollback()
        flash('An error occurred while deleting the form.', 'danger')
        app.logger.error(f"Form deletion error: {str(e)}")

    return redirect(url_for('dashboard'))

@app.route('/update_form/<form_id>', methods=['GET', 'POST'])
@login_required
def update_form(form_id):
    form_response = FormResponse.query.filter_by(form_id=form_id, user_id=current_user.id).first()

    if not form_response:
        flash('Form not found.', 'danger')
        return redirect(url_for('dashboard'))

    form = UpdateFormForm()
    if request.method == 'POST':
        try:
            form_response.form_name = form.form_name.data
            selected_cells = request.form.getlist('selected_cells')
            form_response.selected_cells = ','.join(selected_cells)
            db.session.commit()
            flash('Form updated successfully.', 'success')
            return redirect(url_for('dashboard'))
        except Exception as e:
            db.session.rollback()
            flash('An error occurred while updating the form.', 'danger')
            app.logger.error(f"Form update error: {str(e)}")
    else:
        all_cells = set(form_response.selected_cells.split(','))
        form.selected_cells.choices = [(cell, cell) for cell in all_cells]
        form.selected_cells.data = form_response.selected_cells.split(',')
        form.form_name.data = form_response.form_name
        return render_template('update_form.html', form=form, theme=session.get('theme', 'light'))

@app.errorhandler(404)
def page_not_found(e):
    return render_template('404.html', theme=session.get('theme', 'light')), 404

@app.errorhandler(500)
def internal_error(error):
    return render_template('500.html', theme=session.get('theme', 'light')), 500

def send_email(to, subject, template, **kwargs):
    msg = Message(subject, recipients=[to])
    msg.body = render_template(template + '.txt', **kwargs)
    msg.html = render_template(template + '.html', **kwargs)
    mail.send(msg)

def send_email_reset(to, subject, template, **kwargs):
    msg = Message(subject, recipients=[to])
    msg.body = render_template(template + '.txt', **kwargs)
    msg.html = render_template(template + '_email.html', **kwargs)
    mail.send(msg)

with app.app_context():
    db.create_all()

if __name__ == '__main__':
    app.run(debug=True, port=5001)